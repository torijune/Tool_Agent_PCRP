from agents.table_agents.table_graph.table_workflow_graph import build_table_graph
from agents.table_agents.agent_C.table_parser import load_survey_tables

import pandas as pd
from docx import Document

def get_result(query: str, batch_mode: bool = False, file_path: str = None):
    workflow = build_table_graph()
    state = {
        "query": query,
        "batch_mode": batch_mode,
        "hallucination_reject_num": 0
    }
    if batch_mode:
        state["file_path"] = file_path
    result = workflow.invoke(state)

    hallucination_check = result.get("hallucination_check", "")
    if hallucination_check == "accept":
        return result.get("table_analysis", "⚠️ table_analysis 존재하지 않습니다.")
    elif hallucination_check == "reject":
        return result.get("revised_analysis", "⚠️ revised_analysis 존재하지 않습니다.")
    else:
        return "⚠️ hallucination_check 값이 유효하지 않습니다."

def get_all_result_to_doc(file_path, output_path="analysis_report.docx"):
    tables, question_texts, question_keys = load_survey_tables(file_path)
    workflow = build_table_graph()

    results = []

    print(f"총 {len(question_keys)}개의 질문이 감지되었습니다. Batch 분석을 시작합니다.\n")

    for idx, key in enumerate(question_keys):
        print(f"\n===== [ {key} ] 질문 분석 시작 =====")

        state = {
            "query": f"{question_texts[key]} 분석해줘",
            "file_path": file_path,
            "analysis_type": False,
            "selected_table": tables[key],
            "selected_question": question_texts[key],
            "hallucination_reject_num": 0,
        }

        result = workflow.invoke(state)
        report = result.get("polishing_result", "결과 없음")

        # ✅ 결과 리스트에 저장
        results.append({
            "Index": idx + 1,
            "Question Key": key,
            "Question Text": question_texts[key],
            "Analysis Report": report
        })

    # ✅ Word 문서 생성
    document = Document()
    document.add_heading('Batch 분석 결과 보고서', level=1)

    for item in results:
        document.add_heading(f"Q{item['Index']}: [{item['Question Key']}] {item['Question Text']}", level=2)
        document.add_paragraph(item["Analysis Report"])

    document.save(output_path)
    print(f"\n✅ 분석 결과가 {output_path} 로 저장되었습니다.")


if __name__ == "__main__":
    # # query = input("질문을 입력하세요: ")
    # query = "고양시에서 수행하는 도시 주거 개발 계획에 대해서 궁금해"
    # output = get_result(query)
    # print(output)

    # All Question
    file_path = "/Users/jang-wonjun/Desktop/Dev/Tool_Agent_PCRP/agents/table_agents/table_list/고양시 도시주거환경정비기본계획 조사.xlsx"
    get_all_result_to_doc(file_path, output_path="대기환경_시민인식조사_보고서.docx")